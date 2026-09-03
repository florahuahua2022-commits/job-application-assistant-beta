import unittest
from io import BytesIO
from zipfile import ZipFile

from pypdf import PdfReader

from app.exporter import create_docx, create_pdf, resolve_page_size


SAMPLE = """ALEX MORGAN
Perth, WA | contact@example.com

**PROFESSIONAL SUMMARY**

Experienced administrative professional.

**CORE CAPABILITIES**

- Executive support
- Records management
"""


class ExporterTests(unittest.TestCase):
    def test_docx_is_valid_ooxml(self):
        payload = create_docx(SAMPLE, "Tailored Resume")
        self.assertTrue(payload.startswith(b"PK"))
        with ZipFile(BytesIO(payload)) as archive:
            self.assertIn("word/document.xml", archive.namelist())

    def test_resume_header_is_left_aligned_and_has_compact_margins(self):
        from docx import Document
        document = Document(BytesIO(create_docx(SAMPLE, "Tailored Resume")))
        self.assertEqual(str(document.paragraphs[0].alignment), "LEFT (0)")
        self.assertAlmostEqual(document.sections[0].left_margin.inches, 0.7, places=2)

    def test_pdf_is_readable_and_contains_text(self):
        payload = create_pdf(SAMPLE, "Tailored Resume")
        self.assertTrue(payload.startswith(b"%PDF"))
        reader = PdfReader(BytesIO(payload))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        self.assertIn("ALEX MORGAN", text)
        self.assertIn("Executive support", text)

    def test_modern_docx_uses_arial_theme(self):
        payload = create_docx(SAMPLE, "Tailored Resume", "modern")

        with ZipFile(BytesIO(payload)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
        self.assertIn("Arial", document_xml)

    def test_traditional_pdf_is_readable(self):
        payload = create_pdf(SAMPLE, "Tailored Resume", "traditional")
        reader = PdfReader(BytesIO(payload))

        self.assertIn("ALEX MORGAN", "\n".join(page.extract_text() or "" for page in reader.pages))

    def test_pdf_product_fallback_is_a4_and_date_range_is_ascii(self):
        payload = create_pdf(SAMPLE + "\nJan 2020–Dec 2023", "Tailored Resume")
        reader = PdfReader(BytesIO(payload))
        self.assertAlmostEqual(float(reader.pages[0].mediabox.width), 595.276, places=2)
        self.assertAlmostEqual(float(reader.pages[0].mediabox.height), 841.89, places=2)
        self.assertIn("Jan 2020-Dec 2023", "\n".join(page.extract_text() or "" for page in reader.pages))

    def test_docx_preserves_unicode_and_ascii_date_range(self):
        from docx import Document
        payload = create_docx("José García\ncontact@example.com\nJan 2020—Dec 2023", "Tailored Resume")
        text = "\n".join(paragraph.text for paragraph in Document(BytesIO(payload)).paragraphs)
        self.assertIn("José García", text)
        self.assertIn("Jan 2020-Dec 2023", text)

    def test_market_defaults_and_explicit_override(self):
        expected = {"AU": "A4", "NZ": "A4", "GB": "A4", "UK": "A4", "IE": "A4", "US": "Letter", "CA": "Letter"}
        self.assertEqual({market: resolve_page_size(market)["name"] for market in expected}, expected)
        self.assertEqual(resolve_page_size("AU", "Letter"), {"name": "Letter", "dimensions": (612.0, 792.0), "source": "authoritative_requirement"})
        self.assertEqual(resolve_page_size("unknown")["source"], "product_fallback")

    def test_pdf_market_dimensions_and_templates(self):
        for template in ("classic", "modern", "traditional"):
            a4 = PdfReader(BytesIO(create_pdf(SAMPLE, "Resume", template, "AU"))).pages[0].mediabox
            letter_page = PdfReader(BytesIO(create_pdf(SAMPLE, "Resume", template, "US"))).pages[0].mediabox
            self.assertEqual((round(float(a4.width)), round(float(a4.height))), (595, 842))
            self.assertEqual((round(float(letter_page.width)), round(float(letter_page.height))), (612, 792))

    def test_docx_uses_the_same_resolved_page_size(self):
        from docx import Document
        a4 = Document(BytesIO(create_docx(SAMPLE, "Resume", market="AU"))).sections[0]
        letter_section = Document(BytesIO(create_docx(SAMPLE, "Resume", market="US"))).sections[0]
        self.assertAlmostEqual(a4.page_width.inches, 8.2677, places=3)
        self.assertAlmostEqual(a4.page_height.inches, 11.6929, places=3)
        self.assertAlmostEqual(letter_section.page_width.inches, 8.5, places=3)
        self.assertAlmostEqual(letter_section.page_height.inches, 11, places=3)


if __name__ == "__main__":
    unittest.main()
