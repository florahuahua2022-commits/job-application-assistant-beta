import json
import unittest
from io import BytesIO

from docx import Document

from app.ingest import _extract_scanned_pdf_text, extract_resume_experiences, extract_resume_text, normalise_resume_experiences, parse_job_ad_text, parse_job_page
from app.job_model import build_job_model


class IngestTests(unittest.TestCase):
    def test_ocr_reads_scanned_pdf_pages_and_filters_low_confidence_text(self):
        class FakeImage:
            def close(self):
                pass

        class FakeBitmap:
            def to_pil(self):
                return FakeImage()

            def close(self):
                pass

        class FakePage:
            def render(self, **kwargs):
                self.render_options = kwargs
                return FakeBitmap()

            def close(self):
                pass

        class FakeDocument:
            def __init__(self, payload):
                self.pages = [FakePage(), FakePage()]

            def __len__(self):
                return len(self.pages)

            def __getitem__(self, index):
                return self.pages[index]

            def close(self):
                pass

        class Result:
            txts = ("Alex Morgan", "Project Coordinator", "unreadable")
            scores = (0.99, 0.96, 0.20)

        text = _extract_scanned_pdf_text(
            b"scanned pdf",
            document_factory=FakeDocument,
            ocr_engine=lambda image: Result(),
            image_adapter=lambda image: image,
        )

        self.assertIn("Alex Morgan", text)
        self.assertIn("Project Coordinator", text)
        self.assertNotIn("unreadable", text)

    def test_extracts_master_resume_from_docx(self):
        document = Document()
        document.add_heading("Alex Morgan", level=1)
        document.add_paragraph("Project coordination and administration experience across complex projects.")
        stream = BytesIO()
        document.save(stream)

        text = extract_resume_text("resume.docx", stream.getvalue())

        self.assertIn("Alex Morgan", text)
        self.assertIn("Project coordination", text)

    def test_rejects_unsupported_resume_file(self):
        with self.assertRaises(ValueError):
            extract_resume_text("resume.png", b"not a supported resume document" * 3)

    def test_extracts_structured_experiences_from_resume(self):
        source_text = """
        Alex Morgan
        alex@example.com | 0412 345 678
        Professional Experience
        Project Coordinator
        Bright Energy Pty Ltd
        January 2022 – Present
        Coordinated project schedules, procurement, document control and monthly reporting.
        Prepared stakeholder briefings and maintained accurate project registers.
        Administration Officer | Example Council
        March 2019 - December 2021
        Supported a multidisciplinary team with correspondence, meetings and records management.
        Education
        Bachelor of Business
        """

        result = extract_resume_experiences(source_text)

        self.assertEqual(len(result), 2)
        self.assertEqual(result[0]["role_title"], "Project Coordinator")
        self.assertEqual(result[0]["organization"], "Bright Energy Pty Ltd")
        self.assertIn("project schedules", result[0]["responsibility"])
        self.assertTrue(result[0]["evidence_id"].startswith("EV"))
        self.assertIn("Project Coordinator", result[0]["source_text"])
        self.assertEqual(result[0]["fact_verification"], "explicit")
        self.assertEqual(result[1]["role_title"], "Administration Officer")
        self.assertEqual(result[1]["organization"], "Example Council")

    def test_resume_experience_parser_fails_safely_without_dated_roles(self):
        self.assertEqual(
            extract_resume_experiences("Alex Morgan\nProfessional Summary\nExperienced administrator and coordinator."),
            [],
        )

    def test_employment_header_layout_matrix_preserves_identity_and_explicit_period(self):
        cases = {
            "role_employer_date": "Project Officer\nExample Agency\nFeb 2020 – Present",
            "employer_role_date": "Example Agency\nProject Officer\nFeb 2020 - Current",
            "role_employer_inline_date": "Project Officer\nExample Agency | Feb 2020 – Present",
            "employer_inline_date_role": "Example Agency | Feb 2020 – Present\nProject Officer",
            "all_inline": "Project Officer | Example Agency | Feb 2020 – Present",
            "role_inline_date": "Example Agency\nProject Officer | Feb 2020 – Present",
        }
        for name, header in cases.items():
            with self.subTest(name=name):
                result = extract_resume_experiences(
                    f"Work Experience\n{header}\nPrepared reports, coordinated meetings and maintained accurate project records."
                )
                self.assertEqual(len(result), 1)
                self.assertEqual(result[0]["role_title"], "Project Officer")
                self.assertEqual(result[0]["organization"], "Example Agency")
                self.assertEqual(result[0]["time_period_text"], "Feb 2020 – Present" if "Present" in header else "Feb 2020 - Current")
                self.assertNotRegex(result[0]["organization"], r"20\d{2}")
                self.assertNotRegex(result[0]["responsibility"], r"20\d{2}")

    def test_employer_dashes_are_not_treated_as_header_separators(self):
        for employer in ("Department of Communities – Disability Services", "Department of Communities — Disability Services"):
            with self.subTest(employer=employer):
                result = extract_resume_experiences(
                    f"Work Experience\nFinance Administration Officer\n{employer}\nFeb 2026 – Present\nProvided grounded financial administration and reporting support."
                )
                self.assertEqual(result[0]["organization"], employer)
                self.assertEqual(result[0]["time_period_text"], "Feb 2026 – Present")
        result = extract_resume_experiences(
            "Work Experience\nFinance Administration Officer\nDepartment of Communities – Disability Services | WA State Government\nFeb 2026 – Present\nProvided grounded financial administration and reporting support."
        )
        self.assertEqual(result[0]["organization"], "Department of Communities – Disability Services | WA State Government")

    def test_historical_embedded_periods_are_normalised_without_inference(self):
        historical = [{
            "role_title": "Executive Assistant", "organization": "Avaintec",
            "responsibility": "Nov 2017 – Jan 2019 Prepared agendas and coordinated executive meetings.",
        }, {
            "role_title": "Project Administration Officer",
            "organization": "CCCC Kenya Branch Jan 2016 – Aug 2017",
            "responsibility": "Maintained project records and reporting.",
        }, {
            "role_title": "Project Assistant", "organization": "Pratt & Whitney",
            "responsibility": "Maintained project documentation without a supplied employment date.",
        }]
        result, changed = normalise_resume_experiences(json.dumps(historical))
        experiences = json.loads(result)
        self.assertTrue(changed)
        self.assertEqual(experiences[0]["time_period_text"], "Nov 2017 – Jan 2019")
        self.assertFalse(experiences[0]["responsibility"].startswith("Nov 2017"))
        self.assertEqual(experiences[1]["organization"], "CCCC Kenya Branch")
        self.assertEqual(experiences[1]["time_period_text"], "Jan 2016 – Aug 2017")
        self.assertNotIn("time_period_text", experiences[2])

    def test_reads_structured_job_posting(self):
        html = """
        <html><head><script type="application/ld+json">
        {"@type":"JobPosting","title":"Project Support Officer",
         "hiringOrganization":{"@type":"Organization","name":"Example Energy"},
         "description":"<p>Coordinate projects and prepare reports.</p>"}
        </script></head></html>
        """

        result = parse_job_page(html, "https://example.com/job/1")

        self.assertEqual(result["position_title"], "Project Support Officer")
        self.assertEqual(result["company"], "Example Energy")
        self.assertIn("Coordinate projects", result["job_description"])
        self.assertEqual(result["source"], "structured_job_posting")

    def test_reads_nested_structured_job_posting(self):
        html = """
        <script type="application/ld+json">
        {"page":{"content":{"@type":["Thing","JobPosting"],"title":"Policy Officer",
         "hiringOrganization":{"name":"Example Council"},
         "description":"<p>Prepare policy advice and stakeholder briefings.</p>"}}}
        </script>
        """

        result = parse_job_page(html, "https://example.com/job/2")

        self.assertEqual(result["position_title"], "Policy Officer")
        self.assertEqual(result["company"], "Example Council")

    def test_reads_visible_job_page_when_structured_data_is_missing(self):
        html = """
        <html><head><title>Project Coordinator | SEEK</title></head><body>
        <nav>Sign in</nav><main><h1>Project Coordinator</h1><p>Bright Energy Pty Ltd</p>
        <h2>About the role</h2><p>Coordinate project schedules, procurement, documentation,
        reporting and stakeholder meetings for a growing renewable energy delivery team.</p>
        <h2>What you'll bring</h2><ul><li>Three years of project coordination experience.</li>
        <li>Strong written communication and reporting skills.</li></ul></main></body></html>
        """

        result = parse_job_page(html, "https://example.com/job/3")

        self.assertEqual(result["source"], "page_body")
        self.assertEqual(result["position_title"], "Project Coordinator")
        self.assertEqual(result["company"], "Bright Energy Pty Ltd")
        self.assertIn("project schedules", result["job_description"])

    def test_separates_seek_style_full_job_ad(self):
        raw_text = """
        **Project Support Officer / Junior Project Manager**

        BayWa r.e. Solar Systems Pty Ltd
        Bibra Lake, Perth WA
        Full time

        **About the Role**
        We are seeking a highly organised project support professional to coordinate documentation,
        schedules, stakeholders, procurement and logistics across renewable energy projects.

        **Key Responsibilities**
        Maintain project registers and prepare progress reports for management review.
        """

        result = parse_job_ad_text(raw_text)

        self.assertEqual(result["position_title"], "Project Support Officer / Junior Project Manager")
        self.assertEqual(result["company"], "BayWa r.e. Solar Systems Pty Ltd")
        self.assertEqual(result["warnings"], [])

    def test_warns_when_old_job_content_is_mixed_in(self):
        raw_text = """
        Project Administrator
        Sync Energy
        Perth WA
        About the Role
        Support Fexey battery projects by maintaining schedules, registers and reports.
        About BayWa r.e. Solar Systems
        This older advertisement describes a different solar distribution position and company.
        """

        result = parse_job_ad_text(raw_text, ["BayWa r.e. Solar Systems"])

        self.assertTrue(any("earlier saved job" in warning for warning in result["warnings"]))

    def test_ignores_seek_buttons_and_accepts_western_australia_in_company_name(self):
        raw_text = """
        Program Coordinator
        Minerals Research Institute of Western Australia
        View all jobs
        Share or report ad
        East Perth, Perth WA (Hybrid)
        Full time
        Apply
        Save
        The Position
        Coordinate education and workforce programs, stakeholder relationships, reporting and research scholarships across Western Australia.
        """

        result = parse_job_ad_text(raw_text)

        self.assertEqual(result["company"], "Minerals Research Institute of Western Australia")

    def test_finds_company_from_job_body_when_header_is_missing(self):
        raw_text = """
        Contract Projects Administrators – Perth and Adelaide
        Metrowest is growing and we're looking for capable Projects Administrators to support our project delivery team in a varied and fast-paced role.
        You'll work closely with the project team to keep jobs running smoothly, from purchase orders and invoicing through to document control and reporting.
        What you'll be doing
        Raising purchase orders, maintaining job records, processing invoices and coordinating project documentation.
        What we're looking for
        Strong administration experience in construction, engineering or a similar project environment.
        Why join Metrowest
        Stable, growing business with a long-term project pipeline and supportive team.
        """

        result = parse_job_ad_text(raw_text)

        self.assertEqual(result["company"], "Metrowest")
        self.assertNotEqual(result["company"], "What you'll be doing")

    def test_prefers_labelled_fields_and_extracts_modern_criteria_heading(self):
        raw_text = """
        Job details
        Location
        Perth WA
        Classification
        Administration & Office Support
        Job title: Senior Project Administrator
        Organisation: Horizon Infrastructure
        About the role
        Support complex infrastructure projects through document control, scheduling, reporting,
        procurement coordination and clear communication with internal and external stakeholders.
        What you will bring
        At least three years of project administration experience.
        Advanced document control and Excel skills.
        Strong written communication and attention to detail.
        How to apply
        Submit your resume and cover letter through the application portal.
        """

        result = parse_job_ad_text(raw_text)

        self.assertEqual(result["position_title"], "Senior Project Administrator")
        self.assertEqual(result["company"], "Horizon Infrastructure")
        self.assertEqual(result["selection_criteria"], "")
        model = build_job_model(result["job_description"], result["selection_criteria"])
        self.assertTrue(any("document control" in item["criteria_text"].lower() for item in model["criteria"]))

    def test_srg_private_requirements_do_not_become_selection_criteria_or_include_benefits(self):
        raw_text = """Office Administrator
SRG Global
Perth WA

About the Role
We are looking for an organised and proactive Office Administrator to join our Facades team in Perth CBD.
This is a varied, project-focused administration role supporting procurement, workforce administration, project mobilisation and the day-to-day operation of the Facades office.
You'll work closely with the project team, providing essential administrative support as projects ramp up and workforce requirements increase.

What You'll Be Doing
- Managing office administration systems, processes and documentation.
- Acting as a key point of contact and liaising with internal stakeholders, vendors and service providers.
- Coordinating workforce administration including recruitment support, onboarding, inductions, training, timesheets and compliance records.
- Supporting procurement and project administration including purchase orders, vendor onboarding, PPE procurement and project reporting.
- Maintaining facilities, equipment, supplies, SharePoint and business records.
- Supporting the General Manager, Construction Manager and project teams.
- Contributing to continuous improvement and compliance with systems, safety and quality requirements.

About You
- Previous administration experience, preferably within construction, mining or similar.
- Strong attention to detail and ability to manage competing priorities.
- Excellent written and verbal communication.
- Good time management and organisational skills.
- Microsoft Office and Excel.
- D365 and/or Humanforce highly regarded.

We Offer
- Competitive salary package.
- Career development and progression opportunities.
- Corporate health insurance discounts.
- Corporate discounts on travel, novated leasing and lifestyle benefits.
- Supportive and collaborative team environment."""
        result = parse_job_ad_text(raw_text)
        model = build_job_model(result["job_description"], result["selection_criteria"])
        criteria = " ".join(item["criteria_text"] for item in model["criteria"]).lower()

        self.assertEqual(result["selection_criteria"], "")
        self.assertIn("administration experience", criteria)
        self.assertIn("d365", criteria)
        self.assertNotIn("salary", criteria)
        self.assertNotIn("health insurance", criteria)

    def test_expands_short_zoo_heading_to_readable_advertised_name(self):
        raw_text = """
        Project Coordinator
        ZOO
        Region
        Perth
        The role supports minor works delivered for visitors at Perth Zoo.
        Key responsibilities
        Coordinate project records, contractors and scheduled activities across the site.
        Selection criteria
        Demonstrated project coordination and written communication skills.
        """

        result = parse_job_ad_text(raw_text)

        self.assertEqual(result["company"], "Perth Zoo")


if __name__ == "__main__":
    unittest.main()
