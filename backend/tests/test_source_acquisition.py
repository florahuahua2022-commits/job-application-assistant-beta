import json
import unittest
from email.message import Message
from io import BytesIO
from types import SimpleNamespace
from unittest.mock import patch
from urllib.request import Request
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from reportlab.pdfgen import canvas

from app.ingest import _SafeRedirectHandler, extract_document_text
from app.source_acquisition import MAX_UPLOAD_BYTES, acquire_sources


class FakeResponse:
    def __init__(self, payload: bytes, content_type: str, url="https://files.example/document", filename=None, length=None):
        self.payload = payload
        self.offset = 0
        self.url = url
        self.headers = Message()
        self.headers["Content-Type"] = content_type
        self.headers["Content-Length"] = str(len(payload) if length is None else length)
        if filename:
            self.headers["Content-Disposition"] = f'attachment; filename="{filename}"'

    def read(self, size=-1):
        if self.offset >= len(self.payload):
            return b""
        end = len(self.payload) if size < 0 else min(len(self.payload), self.offset + size)
        chunk, self.offset = self.payload[self.offset:end], end
        return chunk

    def geturl(self):
        return self.url

    def __enter__(self):
        return self

    def __exit__(self, *_):
        pass


class FakeOpener:
    def __init__(self, responses):
        self.responses = list(responses)
        self.requests = []

    def __call__(self, *_):
        return self

    def open(self, request, timeout):
        self.requests.append((request.full_url, timeout))
        return self.responses.pop(0)


def source(source_type="job_description_attachment", confidence="high", number=1):
    return SimpleNamespace(
        source_id=f"source-{number}", source_type=source_type, classification_confidence=confidence,
        source_url=f"https://files.example/{number}", filename=None, content_type=None,
        acquisition_status="discovered", extraction_status="not_attempted", extracted_text="",
        content_sha256=None, warnings_json="[]",
    )


def pdf_payload(pages=1, text="Job description and selection criteria for the advertised position."):
    stream = BytesIO()
    document = canvas.Canvas(stream)
    for _ in range(pages):
        document.drawString(50, 750, text)
        document.showPage()
    document.save()
    return stream.getvalue()


def docx_payload():
    document = Document()
    document.add_paragraph("Position description and application instructions.")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Selection criterion from a table."
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


class SourceAcquisitionTests(unittest.TestCase):
    def acquire(self, sources, responses):
        opener = FakeOpener(responses)
        with patch("app.source_acquisition._validate_public_url", side_effect=lambda value: value):
            acquire_sources(sources, opener)
        return opener

    def test_fetches_pdf_docx_and_plain_text_including_extensionless_pdf(self):
        sources = [source(number=index) for index in range(1, 4)]
        responses = [
            FakeResponse(pdf_payload(), "application/pdf"),
            FakeResponse(docx_payload(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="position.docx"),
            FakeResponse(b"Application instructions in plain text.", "text/plain", filename="instructions.txt"),
        ]

        opener = self.acquire(sources, responses)

        self.assertEqual(len(opener.requests), 3)
        self.assertTrue(all(timeout == 15 for _, timeout in opener.requests))
        self.assertTrue(all(item.acquisition_status == "fetched" for item in sources))
        self.assertTrue(all(item.extraction_status == "extracted" for item in sources))
        self.assertIn("Job description", sources[0].extracted_text)
        self.assertIn("Selection criterion from a table", sources[1].extracted_text)
        self.assertIn("plain text", sources[2].extracted_text)
        self.assertTrue(all(item.content_sha256 for item in sources))

    def test_oversized_response_is_rejected_without_reading_content(self):
        item = source()
        response = FakeResponse(b"not read", "application/pdf", length=MAX_UPLOAD_BYTES + 1)
        self.acquire([item], [response])

        self.assertEqual(item.acquisition_status, "failed")
        self.assertEqual(item.extraction_status, "not_attempted")
        self.assertEqual(response.offset, 0)

    def test_private_redirect_is_rejected_and_redirect_limit_is_five(self):
        with self.assertRaises(ValueError):
            _SafeRedirectHandler().redirect_request(Request("https://example.com"), None, 302, "", {}, "http://127.0.0.1/file.pdf")
        self.assertEqual(_SafeRedirectHandler.max_redirections, 5)

    def test_mime_mismatch_and_login_html_are_not_extracted(self):
        mismatch, login = source(number=1), source(number=2)
        self.acquire([mismatch, login], [
            FakeResponse(docx_payload(), "application/pdf"),
            FakeResponse(b"<html><body>Please sign in to continue</body></html>", "text/html"),
        ])

        self.assertEqual((mismatch.acquisition_status, mismatch.extraction_status), ("unsupported", "not_attempted"))
        self.assertEqual((login.acquisition_status, login.extraction_status), ("requires_auth", "not_attempted"))
        self.assertFalse(mismatch.extracted_text or login.extracted_text)

    def test_corrupt_pdf_is_fetched_but_extraction_fails(self):
        item = source()
        self.acquire([item], [FakeResponse(b"%PDF-corrupt", "application/pdf")])

        self.assertEqual((item.acquisition_status, item.extraction_status), ("fetched", "failed"))
        self.assertTrue(item.content_sha256)
        self.assertEqual(item.extracted_text, "")

    def test_docx_zip_expansion_protection(self):
        stream = BytesIO()
        with ZipFile(stream, "w", ZIP_DEFLATED) as archive:
            for index in range(1001):
                archive.writestr(f"word/item-{index}.xml", "x")
        with self.assertRaisesRegex(ValueError, "safe processing limit"):
            extract_document_text("unsafe.docx", stream.getvalue(), "docx")

    def test_scanned_pdf_uses_existing_ocr_and_long_pdf_is_partial(self):
        blank = pdf_payload(text="")
        with patch("app.ingest._extract_scanned_pdf_text", return_value="OCR extracted selection criterion text") as ocr:
            text, status, _ = extract_document_text("scan.pdf", blank, "pdf")
        self.assertIn("OCR extracted", text)
        self.assertEqual(status, "extracted")
        ocr.assert_called_once()

        _, status, warnings = extract_document_text("long.pdf", pdf_payload(51), "pdf")
        self.assertEqual(status, "partial")
        self.assertIn("first 50", warnings[0])

    def test_low_confidence_and_mandatory_sources_are_not_fetched(self):
        low = source(confidence="low", number=1)
        mandatory = source(source_type="mandatory_form", number=2)
        opener = self.acquire([low, mandatory], [])

        self.assertEqual(opener.requests, [])
        self.assertTrue(all(item.acquisition_status == "discovered" for item in (low, mandatory)))

    def test_maximum_three_downloads_and_twenty_mb_aggregate_budget(self):
        four = [source(number=index) for index in range(4)]
        opener = self.acquire(four, [FakeResponse(b"text attachment", "text/plain", filename="a.txt") for _ in range(3)])
        self.assertEqual(len(opener.requests), 3)
        self.assertEqual(four[3].acquisition_status, "discovered")

        large = [source(number=index) for index in range(3)]
        payload = b"x" * (8 * 1024 * 1024)
        self.acquire(large, [FakeResponse(payload, "text/plain", filename="a.txt") for _ in range(3)])
        self.assertEqual(large[2].acquisition_status, "failed")
        self.assertIn("remaining safe download limit", json.loads(large[2].warnings_json)[0])

    def test_duplicate_content_is_not_extracted_twice(self):
        first, second = source(number=1), source(number=2)
        payload = b"Same application instructions for both links."
        self.acquire([first, second], [
            FakeResponse(payload, "text/plain", filename="one.txt"),
            FakeResponse(payload, "text/plain", filename="two.txt"),
        ])

        self.assertEqual(first.content_sha256, second.content_sha256)
        self.assertEqual(first.extraction_status, "extracted")
        self.assertEqual(second.extraction_status, "not_applicable")
        self.assertEqual(second.extracted_text, "")
        self.assertIn("Duplicate content", json.loads(second.warnings_json)[0])


if __name__ == "__main__":
    unittest.main()
