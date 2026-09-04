import ipaddress
import json
import re
import socket
from html import unescape
from html.parser import HTMLParser
from io import BytesIO
from urllib.parse import unquote, urldefrag, urljoin, urlparse
from urllib.request import HTTPRedirectHandler, Request, build_opener
from zipfile import BadZipFile, ZipFile

from docx import Document
from pypdf import PdfReader
from .ckb import EMPLOYMENT_PERIOD_PATTERN, stable_evidence_id


MAX_UPLOAD_BYTES = 10 * 1024 * 1024
MAX_PAGE_BYTES = 3 * 1024 * 1024
MAX_OCR_PAGES = 6
MAX_ATTACHMENT_PDF_PAGES = 50
MAX_DOCX_ENTRIES = 1000
MAX_DOCX_UNCOMPRESSED_BYTES = 50 * 1024 * 1024


def _extract_scanned_pdf_text(payload: bytes, document_factory=None, ocr_engine=None, image_adapter=None) -> str:
    """OCR a bounded number of PDF pages without sending the resume to another service."""
    if document_factory is None:
        try:
            import pypdfium2 as pdfium
        except ImportError as error:
            raise ValueError("OCR is not available on this server. Try a DOCX or text-based PDF.") from error
        document_factory = pdfium.PdfDocument
    if ocr_engine is None:
        try:
            from rapidocr import RapidOCR
        except ImportError as error:
            raise ValueError("OCR is not available on this server. Try a DOCX or text-based PDF.") from error
        ocr_engine = RapidOCR()
    if image_adapter is None:
        import numpy as np
        image_adapter = np.asarray

    document = document_factory(payload)
    extracted_pages: list[str] = []
    try:
        for page_index in range(min(len(document), MAX_OCR_PAGES)):
            page = document[page_index]
            bitmap = None
            image = None
            try:
                bitmap = page.render(scale=2.0, grayscale=True)
                image = bitmap.to_pil()
                result = ocr_engine(image_adapter(image))
                texts = tuple(getattr(result, "txts", ()) or ())
                scores = tuple(getattr(result, "scores", ()) or ())
                accepted = [
                    str(value).strip()
                    for index, value in enumerate(texts)
                    if str(value).strip() and (index >= len(scores) or float(scores[index]) >= 0.50)
                ]
                if accepted:
                    extracted_pages.append("\n".join(accepted))
            finally:
                if image is not None and hasattr(image, "close"):
                    image.close()
                if bitmap is not None and hasattr(bitmap, "close"):
                    bitmap.close()
                if hasattr(page, "close"):
                    page.close()
    finally:
        if hasattr(document, "close"):
            document.close()
    return "\n\n".join(extracted_pages).strip()


def _resume_line(value: str) -> str:
    value = re.sub(r"^[\s•●▪◦*-]+", "", value.strip())
    return re.sub(r"\s+", " ", value).strip()


_ROLE_HINT = re.compile(r"(?i)\b(?:officer|assistant|administrator|coordinator|manager|director|advisor|adviser|consultant|analyst|specialist|lead|engineer|accountant|clerk|secretary|executive)\b")
_COMPANY_HINT = re.compile(r"(?i)\b(?:pty|ltd|limited|inc|group|services|solutions|council|department|university|college|government|authority|agency|company|corporation|corp|project|branch)\b")


def _employment_identity(first: str, second: str) -> tuple[str, str]:
    first_role, second_role = bool(_ROLE_HINT.search(first)), bool(_ROLE_HINT.search(second))
    if first_role != second_role:
        return (first, second) if first_role else (second, first)
    first_company, second_company = bool(_COMPANY_HINT.search(first)), bool(_COMPANY_HINT.search(second))
    if first_company != second_company:
        return (second, first) if first_company else (first, second)
    return first, second


def normalise_resume_experiences(experiences_json: str) -> tuple[str, bool]:
    try:
        experiences = json.loads(experiences_json or "[]")
    except (TypeError, json.JSONDecodeError):
        return "[]", False
    if not isinstance(experiences, list):
        return "[]", False
    changed = False
    for item in experiences:
        if not isinstance(item, dict) or item.get("time_period_text"):
            continue
        period = item.get("time_period") or {}
        if period.get("start") or period.get("end"):
            continue
        for field, beginning_only in (("organization", False), ("responsibility", True)):
            value = str(item.get(field) or "")
            match = EMPLOYMENT_PERIOD_PATTERN.search(value)
            if not match or (beginning_only and value[:match.start()].strip(" |–—-")):
                continue
            item["time_period_text"] = match.group(0).strip()
            item[field] = _resume_line(f"{value[:match.start()]} {value[match.end():]}").strip(" |–—-")
            changed = True
            break
    return json.dumps(experiences, ensure_ascii=False), changed


def extract_resume_experiences(source_text: str) -> list[dict]:
    """Extract conservative, user-reviewable work-history records from resume text."""
    lines = [_resume_line(line) for line in source_text.splitlines()]
    lines = [line for line in lines if line]
    section_start = next(
        (index + 1 for index, line in enumerate(lines) if re.fullmatch(
            r"(?i)(?:professional |relevant )?(?:work |employment )?(?:experience|history)|employment history|career history",
            line,
        )),
        0,
    )
    section_end = next(
        (index for index in range(section_start, len(lines)) if re.fullmatch(
            r"(?i)(?:education|qualifications|certifications?|skills|technical skills|referees?|references|volunteering)",
            lines[index],
        )),
        len(lines),
    )
    work_lines = lines[section_start:section_end]
    date_indexes = [index for index, line in enumerate(work_lines) if EMPLOYMENT_PERIOD_PATTERN.search(line)]
    if not date_indexes:
        return []

    headers: list[tuple[int, int, str, str, str]] = []
    for date_index in date_indexes:
        line = work_lines[date_index]
        match = EMPLOYMENT_PERIOD_PATTERN.search(line)
        period = match.group(0).strip()
        inline = _resume_line(f"{line[:match.start()]} {line[match.end():]}").strip(" |–—-")
        previous = work_lines[max(0, date_index - 2):date_index]
        responsibility_start = date_index + 1
        if "|" in inline:
            parts = [_resume_line(part) for part in inline.split("|") if _resume_line(part)]
            role_title, organization = _employment_identity(parts[0], parts[1]) if len(parts) >= 2 else ("", "")
            header_start = date_index
        elif inline and _ROLE_HINT.search(inline):
            role_title, organization = inline, previous[-1] if previous else ""
            header_start = date_index - 1 if previous else date_index
        elif inline:
            next_line = work_lines[date_index + 1] if date_index + 1 < len(work_lines) else ""
            if _ROLE_HINT.search(next_line):
                role_title, organization = next_line, inline
                header_start, responsibility_start = date_index, date_index + 2
            else:
                role_title, organization = (previous[-1], inline) if previous else ("", inline)
                header_start = date_index - 1 if previous else date_index
        elif previous and "|" in previous[-1]:
            parts = [_resume_line(part) for part in previous[-1].split("|") if _resume_line(part)]
            if len(previous) >= 2 and _ROLE_HINT.search(previous[-2]) and not any(_ROLE_HINT.search(part) for part in parts):
                role_title, organization, header_start = previous[-2], previous[-1], date_index - 2
            else:
                role_title, organization = _employment_identity(parts[0], parts[1]) if len(parts) >= 2 else ("", "")
                header_start = date_index - 1
        elif len(previous) >= 2:
            role_title, organization = _employment_identity(previous[-2], previous[-1])
            header_start = date_index - 2
        elif previous:
            role_title, organization = previous[-1], ""
            header_start = date_index - 1
        else:
            role_title, organization, header_start = "", "", date_index
        headers.append((header_start, responsibility_start, role_title[:160], organization[:160], period))

    experiences: list[dict] = []
    for position, (header_start, responsibility_start, role_title, organization, period) in enumerate(headers):
        next_header_start = headers[position + 1][0] if position + 1 < len(headers) else len(work_lines)
        responsibility_lines = [
            line for line in work_lines[responsibility_start:next_header_start]
            if len(line) > 2 and not re.fullmatch(r"(?i)(?:responsibilities|key achievements|achievements|duties):?", line)
        ]
        responsibility = " ".join(responsibility_lines).strip()
        if not role_title or len(responsibility) < 25:
            continue
        source_block = "\n".join(work_lines[max(0, header_start):next_header_start]).strip()
        evidence_id = stable_evidence_id("experience", source_block)
        experiences.append({
            "id": evidence_id,
            "evidence_id": evidence_id,
            "evidence_type": "experience",
            "role_title": role_title,
            "organization": organization,
            "responsibility": responsibility,
            "context": f"Employment dates: {period}",
            "result": "",
            "no_result_data": False,
            "source_section": f"Work Experience > {organization or 'Unknown organisation'} > {role_title}",
            "source_text": source_block,
            "time_period_text": period,
            "competency_tags": [],
            "fact_verification": "explicit",
        })
    return experiences[:20]


def extract_document_text(filename: str, payload: bytes, kind: str | None = None) -> tuple[str, str, list[str]]:
    """Extract a bounded PDF, DOCX or text document and report complete/partial status."""
    if not payload:
        raise ValueError("The selected file is empty.")
    if len(payload) > MAX_UPLOAD_BYTES:
        raise ValueError("The file is larger than 10 MB.")
    suffix = kind or (filename.lower().rsplit(".", 1)[-1] if "." in filename else "")
    warnings: list[str] = []
    status = "extracted"
    if suffix == "docx":
        try:
            with ZipFile(BytesIO(payload)) as archive:
                entries = archive.infolist()
                total_size = sum(item.file_size for item in entries)
                if len(entries) > MAX_DOCX_ENTRIES or total_size > MAX_DOCX_UNCOMPRESSED_BYTES:
                    raise ValueError("The DOCX archive expands beyond the safe processing limit.")
                if any(item.file_size > max(1, item.compress_size) * 200 for item in entries):
                    raise ValueError("The DOCX archive has an unsafe compression ratio.")
                if "word/document.xml" not in archive.namelist() or archive.testzip() is not None:
                    raise ValueError("The DOCX file is corrupt or incomplete.")
        except BadZipFile as error:
            raise ValueError("The DOCX file is corrupt or incomplete.") from error
        document = Document(BytesIO(payload))
        text = "\n".join([
            *(paragraph.text for paragraph in document.paragraphs),
            *(cell.text for table in document.tables for row in table.rows for cell in row.cells),
        ])
    elif suffix == "pdf":
        reader = PdfReader(BytesIO(payload), strict=False)
        page_count = len(reader.pages)
        text = "\n".join(reader.pages[index].extract_text() or "" for index in range(min(page_count, MAX_ATTACHMENT_PDF_PAGES)))
        if page_count > MAX_ATTACHMENT_PDF_PAGES:
            status = "partial"
            warnings.append(f"Only the first {MAX_ATTACHMENT_PDF_PAGES} PDF pages were processed.")
        if len(re.sub(r"\s+", "", text)) < 40:
            text = _extract_scanned_pdf_text(payload)
            if page_count > MAX_OCR_PAGES:
                status = "partial"
                warnings.append(f"OCR was limited to the first {MAX_OCR_PAGES} PDF pages.")
    elif suffix in {"txt", "md"}:
        text = payload.decode("utf-8-sig", errors="replace")
    else:
        raise ValueError("Only DOCX, PDF and plain-text files are supported.")
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    return text, status, list(dict.fromkeys(warnings))


def extract_resume_text(filename: str, payload: bytes) -> str:
    try:
        text, _, _ = extract_document_text(filename, payload)
    except ValueError as error:
        if str(error) == "The selected file is empty.":
            raise ValueError("The selected resume file is empty.") from error
        if str(error) == "The file is larger than 10 MB.":
            raise ValueError("The resume file is larger than 10 MB.") from error
        if "Only DOCX" in str(error):
            raise ValueError("Upload a DOCX, PDF or TXT resume file.") from error
        raise
    if len(text) < 40:
        raise ValueError("Very little text could be read from this file. Try a DOCX or text-based PDF.")
    return text


class _JobPageParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.meta: dict[str, str] = {}
        self.json_ld: list[str] = []
        self._in_json_ld = False
        self._json_parts: list[str] = []
        self._hidden_depth = 0
        self._in_title = False
        self.title_parts: list[str] = []
        self.body_parts: list[str] = []
        self.links: list[dict] = []
        self._active_link: dict | None = None
        self._contexts: list[tuple[str, list[str]]] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        tag = tag.lower()
        values = {key.lower(): value or "" for key, value in attrs}
        if tag == "meta":
            key = (values.get("property") or values.get("name")).lower()
            if key and values.get("content"):
                self.meta[key] = values["content"]
        if tag == "script" and "ld+json" in values.get("type", "").lower():
            self._in_json_ld = True
            self._json_parts = []
        elif tag in {"script", "style", "noscript", "svg"}:
            self._hidden_depth += 1
        if tag == "title":
            self._in_title = True
        if not self._hidden_depth and tag in {"p", "li"}:
            self._contexts.append((tag, []))
        if not self._hidden_depth and tag == "a" and values.get("href"):
            self._active_link = {"href": values["href"], "title": values.get("title", ""), "label_parts": [], "context_parts": self._contexts[-1][1] if self._contexts else []}
        if tag in {"p", "li", "div", "section", "article", "h1", "h2", "h3", "br"}:
            self.body_parts.append("\n")

    def handle_data(self, data: str) -> None:
        if self._in_json_ld:
            self._json_parts.append(data)
        elif not self._hidden_depth:
            if self._contexts:
                self._contexts[-1][1].append(data)
            if self._active_link is not None:
                self._active_link["label_parts"].append(data)
            if self._in_title:
                self.title_parts.append(data)
            else:
                self.body_parts.append(data)

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag == "script" and self._in_json_ld:
            self.json_ld.append("".join(self._json_parts))
            self._in_json_ld = False
        elif tag in {"script", "style", "noscript", "svg"} and self._hidden_depth:
            self._hidden_depth -= 1
        if tag == "a" and self._active_link is not None:
            self.links.append(self._active_link)
            self._active_link = None
        if tag == "title":
            self._in_title = False
        if tag in {"p", "li"} and self._contexts and self._contexts[-1][0] == tag:
            self._contexts.pop()
        if tag in {"p", "li", "div", "section", "article", "h1", "h2", "h3"}:
            self.body_parts.append("\n")

    def discovered_links(self, page_url: str) -> list[dict]:
        result: list[dict] = []
        seen: set[str] = set()
        for link in self.links:
            resolved, _ = urldefrag(urljoin(page_url, link["href"]))
            parsed = urlparse(resolved)
            if parsed.scheme not in {"http", "https"} or not parsed.netloc or resolved in seen:
                continue
            seen.add(resolved)
            result.append({
                "url": resolved,
                "href": link["href"],
                "label": re.sub(r"\s+", " ", "".join(link["label_parts"])).strip(),
                "title": re.sub(r"\s+", " ", link["title"]).strip(),
                "context": re.sub(r"\s+", " ", "".join(link["context_parts"])).strip()[:1000],
                "filename": unquote(parsed.path.rsplit("/", 1)[-1]),
                "discovered_from_url": page_url,
            })
        return result


def _clean_html(value: str) -> str:
    value = re.sub(r"(?i)<br\s*/?>", "\n", value)
    value = re.sub(r"(?i)</(?:p|li|h[1-6]|div)>", "\n", value)
    value = re.sub(r"<[^>]+>", "", value)
    value = unescape(value).replace("\xa0", " ")
    value = re.sub(r"[ \t]+", " ", value)
    return re.sub(r"\n{3,}", "\n\n", value).strip()


def _job_postings(value):
    if isinstance(value, list):
        for item in value:
            yield from _job_postings(item)
    elif isinstance(value, dict):
        item_type = value.get("@type", [])
        item_types = [item_type] if isinstance(item_type, str) else item_type
        if "JobPosting" in item_types:
            yield value
        for child in value.values():
            if isinstance(child, (dict, list)):
                yield from _job_postings(child)


def parse_job_page(html: str, url: str) -> dict:
    parser = _JobPageParser()
    parser.feed(html)
    discovered_sources = parser.discovered_links(url)
    for raw in parser.json_ld:
        try:
            candidates = list(_job_postings(json.loads(raw)))
        except (json.JSONDecodeError, TypeError):
            continue
        if candidates:
            job = candidates[0]
            organisation = job.get("hiringOrganization") or {}
            company = organisation.get("name", "") if isinstance(organisation, dict) else str(organisation)
            return {
                "company": _clean_html(company),
                "position_title": _clean_html(str(job.get("title", ""))),
                "job_description": _clean_html(str(job.get("description", ""))),
                "job_url": url,
                "source": "structured_job_posting",
                "discovered_sources": discovered_sources,
            }

    page_body = _clean_html("".join(parser.body_parts))
    title = parser.meta.get("og:title") or " ".join(parser.title_parts)
    description = parser.meta.get("og:description") or parser.meta.get("description", "")
    title = re.sub(r"\s*[|\-–—]\s*(SEEK|Indeed|LinkedIn|Jora|Glassdoor).*", "", title, flags=re.I).strip()
    if len(page_body) >= 120:
        try:
            parsed_body = parse_job_ad_text(page_body)
        except ValueError:
            parsed_body = {}
        if parsed_body:
            return {
                "company": parsed_body.get("company", ""),
                "position_title": parsed_body.get("position_title") or _clean_html(title),
                "job_description": parsed_body.get("job_description") or _clean_html(description),
                "job_url": url,
                "source": "page_body",
                "discovered_sources": discovered_sources,
            }
    return {
        "company": "",
        "position_title": _clean_html(title),
        "job_description": _clean_html(description),
        "job_url": url,
        "source": "page_summary",
        "discovered_sources": discovered_sources,
    }


def _plain_ad_line(value: str) -> str:
    value = re.sub(r"!\[[^]]*\]\([^)]*\)", "", value)
    value = re.sub(r"\[([^]]+)\]\([^)]*\)", r"\1", value)
    value = value.replace("**", "").replace("__", "").strip(" #*\t")
    return re.sub(r"\s+", " ", value).strip()


def expand_abbreviated_company(company: str, text: str) -> str:
    """Prefer a readable advertised name when extraction returns a short all-caps fragment."""
    if not re.fullmatch(r"[A-Z]{2,5}", company.strip()):
        return company
    suffix = company.strip().title()
    matches = re.findall(rf"\b(?:[A-Z][a-z]+\s+){{1,3}}{re.escape(suffix)}\b", text)
    excluded_prefixes = {"region", "location", "organisation", "organization", "employer"}
    candidates = [
        _plain_ad_line(match) for match in matches
        if _plain_ad_line(match).split()[0].lower() not in excluded_prefixes
    ]
    return min(candidates, key=lambda value: (len(value.split()), len(value)), default=company)


def parse_job_ad_text(raw_text: str, previous_companies: list[str] | None = None) -> dict:
    text = raw_text.strip()
    if len(text) < 120:
        raise ValueError("Paste the complete job advertisement, not only the title or link.")
    lines = [_plain_ad_line(line) for line in text.splitlines()]
    lines = [line for line in lines if line]
    noise = re.compile(
        r"(?i)^(view all jobs|share or report ad|apply(?: now)?$|save$|posted\b|high application volume|how you match|show all|"
        r"sign in|create (?:a )?job alert|job details$|classification$|subclassification$|location$|work type$|"
        r"salary(?:\s|$)|full[ -]?time$|part[ -]?time$|contract/temp$|employer questions?)"
    )
    candidates = [line for line in lines[:30] if not noise.search(line) and not line.lower().startswith("http")]
    labelled_title = re.search(r"(?im)^\s*(?:job title|position title|position|role)\s*:\s*(.+?)\s*$", text)
    labelled_company = re.search(r"(?im)^\s*(?:company|organisation|organization|employer)\s*:\s*(.+?)\s*$", text)
    position_title = _plain_ad_line(labelled_title.group(1))[:160] if labelled_title else ""
    company = _plain_ad_line(labelled_company.group(1))[:160] if labelled_company else ""
    excluded_heading = re.compile(
        r"(?i)^(about\b|what\b|who we\b|we offer\b|job summary\b|job description\b|key responsibilities\b|"
        r"responsibilities\b|requirements\b|selection criteria\b|the position\b|how to apply\b)"
    )
    if not position_title:
        for candidate in candidates:
            if excluded_heading.search(candidate):
                break  # Job identity belongs above the body headings, not inside the duties.
            if ":" in candidate[:30]:
                continue
            if re.search(r"(?i),\s*(?:perth\s+)?WA(?:\s|\(|$)", candidate):
                continue
            position_title = candidate[:160]
            break
    if not labelled_title and (not position_title or len(position_title.split()) > 12 or re.match(r"(?i)^(?:we|our|you|by)\b", position_title)):
        role_match = re.search(
            r"(?i)\b(?:demand for|seeking|looking for|hiring)[ \t]+(?:an?[ \t]+)?(?:experienced[ \t]+)?"
            r"([A-Z][A-Za-z&/' -]{2,70}?)(?=\s+(?:for|to|who|with|across|in)\b|[.,;]|$)",
            text,
        )
        position_title = _plain_ad_line(role_match.group(1)) if role_match else ""
    company_patterns = (
        r"(?im)^\s*([A-Z][A-Za-z0-9&.'’ -]{2,100}?)\s+(?:is|are)\s+(?:growing|seeking|looking|hiring)\b",
        r"(?im)^\s*why\s+join\s+([A-Z][A-Za-z0-9&.'’ -]{2,100}?)\s*$",
        r"(?i)\b(?:with|at)\s+([A-Z][A-Za-z0-9&.'’ -]{2,80}?)(?=\s*[,.;]|\s+(?:you|we|our|for|to|as)\b)",
    )
    for pattern in company_patterns if not company else ():
        for match in re.finditer(pattern, text):
            candidate = _plain_ad_line(match.group(1))
            if not excluded_heading.search(candidate) and not re.match(r"(?i)^(?:we|our|you|your)\b", candidate):
                company = candidate
                break
        if company:
            break
    for line in candidates[1:8] if not company and candidates and not excluded_heading.search(candidates[0]) else []:
        lowered = line.lower()
        if excluded_heading.search(line) or lowered.startswith(("the role", "location", "why join")):
            break
        is_location = bool(re.search(r"(?i),\s*(?:perth\s+)?WA(?:\s|\(|$)", line))
        is_category = bool(re.search(r"\([^)]*(?:construction|technology|administration|management)[^)]*\)", line, re.I))
        if len(line) <= 140 and len(line.split()) <= 10 and not is_location and not is_category and not excluded_heading.search(line):
            company = line
            break

    company = expand_abbreviated_company(company, text)

    criteria = ""
    criteria_match = re.search(
        r"(?is)(?:key selection criteria|selection criteria|essential criteria)\s*[:\n]+(.+?)"
        r"(?=\n\s*(?:who we are|about us|about (?:the|our) (?:company|organisation|organization)|"
        r"our values|company values|our culture|company culture|we offer|what we offer|benefits|employee benefits|"
        r"our benefits|perks|rewards and benefits|why join us|what(?:'|’)?s in it for you|"
        r"what you(?:'|’)?ll get|how to apply|employer questions?)\b|\Z)",
        text,
    )
    if criteria_match:
        criteria = _clean_html(criteria_match.group(1))

    warnings: list[str] = []
    if len(re.findall(r"(?im)^\s*\**about the role\**\s*$", text)) > 1:
        warnings.append("More than one 'About the Role' section was detected. The text may contain multiple job advertisements.")
    if len(re.findall(r"(?im)^\s*\**employer questions?\**\s*$", text)) > 1:
        warnings.append("More than one employer-question section was detected. Check whether two advertisements were pasted together.")
    lowered_text = text.lower()
    for previous_company in previous_companies or []:
        name = previous_company.strip()
        if len(name) >= 4 and name.lower() in lowered_text and name.lower() not in company.lower():
            warnings.append(f"The text mentions a company from an earlier saved job: {name}. Remove any old job content before saving.")
    if not company:
        warnings.append("The organisation could not be identified confidently. Check it before saving.")
    if not position_title:
        warnings.append("The position title could not be identified confidently. Check it before saving.")

    return {
        "company": company,
        "position_title": position_title,
        "job_description": "\n".join(lines),
        "selection_criteria": criteria,
        "warnings": list(dict.fromkeys(warnings)),
    }


def _validate_public_url(url: str) -> str:
    parsed = urlparse(url.strip())
    if parsed.scheme not in {"http", "https"} or not parsed.hostname:
        raise ValueError("Enter a valid public http or https job link.")
    try:
        addresses = socket.getaddrinfo(parsed.hostname, parsed.port or (443 if parsed.scheme == "https" else 80))
    except socket.gaierror as error:
        raise ValueError("The job website could not be found.") from error
    for address in addresses:
        ip = ipaddress.ip_address(address[4][0])
        if not ip.is_global:
            raise ValueError("Only public job website links can be imported.")
    return parsed.geturl()


class _SafeRedirectHandler(HTTPRedirectHandler):
    max_redirections = 5

    def redirect_request(self, req, fp, code, msg, headers, newurl):
        _validate_public_url(newurl)
        return super().redirect_request(req, fp, code, msg, headers, newurl)


def import_job_url(url: str) -> dict:
    safe_url = _validate_public_url(url)
    request = Request(safe_url, headers={"User-Agent": "Mozilla/5.0 JobApplicationAssistant/1.0"})
    try:
        with build_opener(_SafeRedirectHandler()).open(request, timeout=15) as response:
            final_url = _validate_public_url(response.geturl())
            payload = response.read(MAX_PAGE_BYTES + 1)
            if len(payload) > MAX_PAGE_BYTES:
                raise ValueError("The job page is too large to import safely.")
            charset = response.headers.get_content_charset() or "utf-8"
    except ValueError:
        raise
    except Exception as error:
        raise ValueError("This website did not allow automatic reading. Paste the job details manually instead.") from error
    result = parse_job_page(payload.decode(charset, errors="replace"), final_url)
    if not result["position_title"] and not result["job_description"]:
        raise ValueError("No readable job details were found on this page. Paste the job details manually instead.")
    return result
