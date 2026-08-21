from __future__ import annotations

from io import BytesIO
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


INK = RGBColor(23, 36, 58)
ACCENT = RGBColor(8, 117, 101)
MUTED = RGBColor(95, 107, 122)

EXPORT_THEMES = {
    "classic": {"font": "Calibri", "body_size": 11, "margin": 1.0, "line_spacing": 1.10, "accent": ACCENT, "pdf_font": "Helvetica", "pdf_bold": "Helvetica-Bold"},
    "modern": {"font": "Arial", "body_size": 10.5, "margin": 0.85, "line_spacing": 1.08, "accent": RGBColor(37, 99, 135), "pdf_font": "Helvetica", "pdf_bold": "Helvetica-Bold"},
    "traditional": {"font": "Georgia", "body_size": 11, "margin": 1.0, "line_spacing": 1.12, "accent": RGBColor(52, 63, 82), "pdf_font": "Times-Roman", "pdf_bold": "Times-Bold"},
}

PAGE_SIZES = {"A4": A4, "Letter": letter}
MARKET_PAGE_SIZES = {
    "AU": "A4", "AUSTRALIA": "A4", "NZ": "A4", "NEW ZEALAND": "A4",
    "GB": "A4", "UK": "A4", "UNITED KINGDOM": "A4", "IE": "A4", "IRELAND": "A4",
    "US": "Letter", "USA": "Letter", "UNITED STATES": "Letter",
    "CA": "Letter", "CANADA": "Letter",
}
PRODUCT_PAGE_SIZE = "A4"


def resolve_page_size(market: str | None = None, explicit: str | None = None) -> dict:
    if explicit in PAGE_SIZES:
        name, source = explicit, "authoritative_requirement"
    elif str(market or "").strip().upper() in MARKET_PAGE_SIZES:
        name, source = MARKET_PAGE_SIZES[str(market).strip().upper()], "market_default"
    else:
        name, source = PRODUCT_PAGE_SIZE, "product_fallback"
    return {"name": name, "dimensions": PAGE_SIZES[name], "source": source}


def export_theme(template: str) -> dict:
    return EXPORT_THEMES.get(template, EXPORT_THEMES["classic"])


def safe_filename(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", value.strip())
    return cleaned.strip("_") or "application_document"


def _ascii_punctuation(text: str) -> str:
    return text.translate(str.maketrans({
        "–": "-", "—": "-", "‑": "-", "“": '"', "”": '"',
        "‘": "'", "’": "'", "…": "...", "•": "-",
    }))


def _set_font(run, size: float = 11, bold: bool = False, color: RGBColor = INK, font_name: str = "Calibri") -> None:
    run.font.name = font_name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def _add_page_number(paragraph, theme: dict) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    _set_font(run, size=9, color=MUTED, font_name=theme["font"])
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText"); instruction.set(qn("xml:space"), "preserve"); instruction.text = " PAGE "
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def _configure_docx(document: Document, theme: dict, page_size: tuple[float, float]) -> None:
    section = document.sections[0]
    section.page_width = Inches(page_size[0] / 72)
    section.page_height = Inches(page_size[1] / 72)
    section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(theme["margin"])
    section.header_distance = section.footer_distance = Inches(0.492)
    _add_page_number(section.footer.paragraphs[0], theme)

    normal = document.styles["Normal"]
    normal.font.name = theme["font"]; normal.font.size = Pt(theme["body_size"]); normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = theme["line_spacing"]
    for style_name, size, color, before, after in (
        ("Heading 1", 16, ACCENT, 16, 8),
        ("Heading 2", 13, ACCENT, 12, 6),
        ("Heading 3", 12, INK, 8, 4),
    ):
        style = document.styles[style_name]
        style.font.name = theme["font"]; style.font.size = Pt(size); style.font.bold = True; style.font.color.rgb = theme["accent"] if color == ACCENT else color
        style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def _add_inline_markdown(paragraph, text: str, theme: dict) -> None:
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        run = paragraph.add_run(part[2:-2] if bold else part)
        _set_font(run, size=theme["body_size"], bold=bold, font_name=theme["font"])


def create_docx(content: str, title: str, template: str = "classic", market: str | None = None, page_size: str | None = None) -> bytes:
    theme = export_theme(template)
    document = Document()
    _configure_docx(document, theme, resolve_page_size(market, page_size)["dimensions"])
    lines = _ascii_punctuation(content).splitlines()

    for index, raw in enumerate(lines):
        line = raw.strip()
        if not line:
            document.add_paragraph()
            continue
        if line.startswith("### "):
            document.add_paragraph(line[4:], style="Heading 3")
        elif line.startswith("## "):
            document.add_paragraph(line[3:], style="Heading 2")
        elif line.startswith("# "):
            document.add_paragraph(line[2:], style="Heading 1")
        elif line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            paragraph.paragraph_format.space_after = Pt(4)
            _add_inline_markdown(paragraph, line[2:], theme)
        elif re.match(r"^\d+\.\s", line):
            paragraph = document.add_paragraph(style="List Number")
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            paragraph.paragraph_format.space_after = Pt(4)
            _add_inline_markdown(paragraph, re.sub(r"^\d+\.\s*", "", line), theme)
        elif line.startswith("**") and line.endswith("**") and len(line) < 100:
            paragraph = document.add_paragraph(style="Heading 2")
            paragraph.add_run(line[2:-2])
        elif index == 0 and len(line) < 100:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(5)
            run = paragraph.add_run(line.replace("**", "")); _set_font(run, size=19, bold=True, color=theme["accent"], font_name=theme["font"])
        elif index == 1 and len(line) < 140:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(14)
            run = paragraph.add_run(line.replace("**", "")); _set_font(run, size=10, color=MUTED, font_name=theme["font"])
        else:
            paragraph = document.add_paragraph()
            _add_inline_markdown(paragraph, line, theme)

    stream = BytesIO(); document.save(stream)
    return stream.getvalue()


def create_pdf(content: str, title: str, template: str = "classic", market: str | None = None, page_size: str | None = None) -> bytes:
    theme = export_theme(template)
    accent_hex = "#" + "".join(f"{channel:02X}" for channel in theme["accent"])
    stream = BytesIO()
    styles = getSampleStyleSheet()
    body = ParagraphStyle("ApplicationBody", parent=styles["BodyText"], fontName=theme["pdf_font"], fontSize=theme["body_size"] - .5, leading=(theme["body_size"] + 3), textColor="#17243A", spaceAfter=7)
    heading = ParagraphStyle("ApplicationHeading", parent=body, fontName=theme["pdf_bold"], fontSize=13, leading=16, textColor=accent_hex, spaceBefore=10, spaceAfter=6, keepWithNext=True)
    title_style = ParagraphStyle("ApplicationTitle", parent=heading, fontSize=18, leading=22, alignment=TA_CENTER, spaceAfter=10)
    bullet = ParagraphStyle("ApplicationBullet", parent=body, leftIndent=18, firstLineIndent=-10, bulletIndent=4, spaceAfter=4)
    story = []
    for index, raw in enumerate(_ascii_punctuation(content).splitlines()):
        line = raw.strip()
        if not line:
            story.append(Spacer(1, 5)); continue
        escaped = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        escaped = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", escaped)
        if line.startswith("#"):
            escaped = re.sub(r"^#+\s*", "", escaped)
            story.append(Paragraph(escaped, heading))
        elif line.startswith("- "):
            story.append(Paragraph(escaped[2:], bullet, bulletText="-"))
        elif index == 0 and len(line) < 100:
            story.append(Paragraph(escaped.replace("**", ""), title_style))
        elif line.startswith("**") and line.endswith("**") and len(line) < 100:
            story.append(Paragraph(escaped, heading))
        else:
            story.append(Paragraph(escaped, body))

    margin = theme["margin"] * inch
    pdf = SimpleDocTemplate(stream, pagesize=resolve_page_size(market, page_size)["dimensions"], rightMargin=margin, leftMargin=margin, topMargin=margin, bottomMargin=margin, title=title, author="")
    pdf.build(story)
    return stream.getvalue()
